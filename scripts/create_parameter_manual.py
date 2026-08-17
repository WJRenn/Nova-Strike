from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NovaStrike_参数手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "666666"
HEADER_FILL = "E8EEF5"
NOTE_FILL = "FFF4CE"
RISK_FILL = "FDE9E7"
GREEN_FILL = "E7F3E7"
WHITE = "FFFFFF"
GRID = "AAB7C4"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "5")
        tag.set(qn("w:color"), GRID)


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def set_run_font(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_cell_text(cell, bold=False, color="000000", align=None, size=9.5):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        if align is not None:
            p.alignment = align
        for r in p.runs:
            set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, center_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        style_cell_text(table.rows[0].cells[i], bold=True, color=NAVY,
                        align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            style_cell_text(cells[i], align=align, size=9.3)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, label, text, fill=NOTE_FILL):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    set_cell_shading(cell, fill)
    set_table_geometry(table, [TABLE_WIDTH])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    r = paragraph.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Parameter Label" not in [s.name for s in styles]:
        st = styles.add_style("Parameter Label", WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles["Normal"]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    header = section.header.paragraphs[0]
    header.text = "NOVA STRIKE  |  参数参考手册"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in header.runs:
        set_run_font(r, size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])
    return doc


def build():
    doc = setup_document()

    # Editorial cover pattern.
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOURCE-OF-TRUTH REFERENCE")
    set_run_font(r, size=10, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nova Strike 参数手册")
    set_run_font(r, size=28, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("玩家 · 敌人 · 道具 · 关卡 · 战斗与系统参数")
    set_run_font(r, size=14, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("依据当前工作区源码整理")
    set_run_font(r, size=10.5, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基线日期：2026-08-14  |  C++17 / SFML 3.x")
    set_run_font(r, size=10.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "1. 文档说明", 1)
    add_body(doc, "本手册记录当前源码中实际生效的游戏参数，适合作为后续修改、平衡调整和面试讲解的基准。参数来源为 src 目录下的当前 C++/HPP 文件，而不是 README 中的概括描述。")
    add_callout(doc, "读取规则", "表格中的“理论值”按代码直接推导；涉及随机生成时给出概率、范围或期望值，不把随机结果误写成固定配置。", GREEN_FILL)
    add_table(doc, ["类别", "当前值", "主要源码位置"], [
        ("窗口", "1000 × 1000 像素", "config/GameConfig.hpp"),
        ("目标逻辑更新率", "60 Hz（固定步长 1/60 秒）", "app/Game.hpp / Game.cpp"),
        ("渲染帧率上限", "60 FPS", "app/Game.cpp"),
        ("初始游戏状态", "Menu", "app/Game.cpp"),
        ("关卡阶段数", "5 个可战斗阶段 + Victory 状态", "gameplay/LevelManager.hpp"),
        ("普通敌人类型", "Alpha / Beta / Gamma", "InvaderTypes.hpp"),
        ("Boss 类型", "Dragon / Monster", "Dragon.hpp / Monster.hpp"),
        ("道具/掉落类型", "HealthPickup / ScorePickup / PiercingPickup / MeteorHazard", "src/entities/pickups"),
    ], [1800, 3000, 4560], center_cols=(0,))

    add_heading(doc, "2. 玩家参数", 1)
    add_table(doc, ["参数", "当前值", "说明"], [
        ("初始生命", "3", "创建新玩家时的 health"),
        ("声明最大生命", "3", "MAX_HEALTH；读取存档时会限制到 3"),
        ("游戏内最大生命", "3", "addHealth() 与存档读取共用 MAX_HEALTH"),
        ("初始分数", "0", "击杀敌人或拾取分数道具后增加"),
        ("最大移动速度", "400 像素/秒", "速度向量超过该值时归一化限制"),
        ("加速度", "800 像素/秒²", "按住移动键时增加速度"),
        ("摩擦减速度", "600 像素/秒²", "无方向输入时逐渐减速"),
        ("初始位置", "(450, 850)", "由 1000/2-50 与 1000-150 计算"),
        ("飞船缩放", "0.75 × 0.75", "三种飞船共用"),
        ("精灵动画帧", "131 × 131 像素", "按 8 个方向选择精灵表区域"),
        ("受炸弹伤害", "每次 -1 HP", "普通敌人及 Boss 炸弹相同"),
        ("屏幕边界", "环绕", "越过一侧后从另一侧出现"),
    ], [2300, 2200, 4860], center_cols=(0, 1))
    add_callout(doc, "生命上限", "玩家初始生命、HealthPickup、HUD 和存档读取统一使用最大生命 3。", GREEN_FILL)

    add_heading(doc, "2.1 飞船与操作", 2)
    add_table(doc, ["项目", "键盘模式", "鼠标模式"], [
        ("移动", "WASD 或方向键", "飞船朝鼠标位置移动"),
        ("射击", "Space", "鼠标左键"),
        ("暂停/恢复", "Escape", "Escape"),
        ("飞船类型", "1 红色 / 2 绿色 / 3 橙色", "只改变纹理，不改变属性"),
    ], [1900, 3730, 3730], center_cols=(0,))

    add_heading(doc, "3. 玩家武器与投射物", 1)
    add_table(doc, ["参数", "普通射击", "PiercingPickup 强化射击"], [
        ("每次生成数量", "2 发", "1 发"),
        ("射击冷却", "200 ms", "200 ms（没有实际射速提升）"),
        ("理论最大射速", "5 次射击/秒，即 10 发/秒", "5 发/秒"),
        ("单发伤害", "1 HP", "1 HP"),
        ("是否穿透", "否，命中第一个敌人后销毁", "是，可在同一帧命中多个敌人"),
        ("飞行速度", "600 像素/秒", "600 像素/秒"),
        ("方向", "垂直向上 (0, -1)", "垂直向上 (0, -1)"),
        ("缩放", "0.75", "0.20"),
        ("销毁范围", "y < -50 或 y > 1050", "相同"),
    ], [2200, 3550, 3610], center_cols=(0, 1, 2))
    add_callout(doc, "实现备注", "PiercingPickup 提供 5 秒穿透效果，不改变 200ms 射击冷却。强化弹出生位置为玩家位置 (-15, -180) 偏移，明显高于普通子弹。", NOTE_FILL)

    add_heading(doc, "4. 普通敌人参数", 1)
    add_body(doc, "三种普通敌人均使用 Invaders.png 精灵表，生命耗尽时销毁。每发玩家子弹造成 1 点伤害，因此“击杀所需命中数”等于敌人生命值。")
    add_table(doc, ["类型", "代码", "HP", "分数", "攻击间隔", "弹速", "攻击方式", "移动速度"], [
        ("Alpha", "A", "3", "10", "5 秒", "200", "垂直向下", "30"),
        ("Beta", "B", "5", "20", "3 秒", "350", "-30°~+30°随机散射", "60"),
        ("Gamma", "G", "7", "30", "2 秒", "450", "发射瞬间瞄准玩家", "100"),
    ], [1150, 650, 650, 700, 1100, 850, 2750, 1510], center_cols=(0, 1, 2, 3, 4, 5, 7))
    add_table(doc, ["公共参数", "当前值", "说明"], [
        ("水平摆动范围", "中心点左右各 50 像素", "position.x = startX + sin(timer) × 50"),
        ("摆动计时增长", "dt × 移动速度 × 0.02", "移动速度越高，摆动频率越快"),
        ("同时活动炸弹数", "最多 1 颗", "现有炸弹销毁后才能再次生成"),
        ("炸弹基础伤害", "1 HP", "命中玩家后炸弹销毁"),
        ("炸弹默认出生偏移", "敌人位置 (+40, +50)", "以敌人 position 为基准"),
        ("炸弹销毁边界", "x/y 超出 -50~1050", "适用于斜向弹"),
    ], [2200, 2600, 4560], center_cols=(0, 1))

    add_heading(doc, "5. Boss 参数", 1)
    add_table(doc, ["参数", "Dragon", "Monster"], [
        ("类型代码", "D", "M"),
        ("生命值", "100 HP", "100 HP"),
        ("击杀分数", "100", "100"),
        ("击杀所需普通命中", "100 次", "100 次"),
        ("攻击间隔", "1.5 秒", "1.5 秒"),
        ("炸弹速度", "350 像素/秒", "450 像素/秒"),
        ("每轮弹数", "5 颗", "3 颗"),
        ("攻击方式", "五向散射", "发射瞬间瞄准 + 三向偏移"),
        ("初始关卡位置", "(500, 200)", "(500, 200)"),
        ("水平运动", "中心左右各 250 像素", "无水平摆动"),
        ("垂直运动", "中心上下各 80 像素", "中心上下各 50 像素"),
        ("精灵缩放", "0.30", "0.40"),
    ], [2400, 3480, 3480], center_cols=(0, 1, 2))

    add_heading(doc, "5.1 Dragon 攻击细节", 2)
    add_table(doc, ["弹道索引", "方向向量 (x, y)", "出生 X 偏移", "效果"], [
        ("1", "(-0.6, 1.0)", "-24", "最左侧"),
        ("2", "(-0.3, 1.0)", "-12", "左侧"),
        ("3", "(0.0, 1.0)", "0", "正下方"),
        ("4", "(0.3, 1.0)", "+12", "右侧"),
        ("5", "(0.6, 1.0)", "+24", "最右侧"),
    ], [1400, 2500, 2100, 3360], center_cols=(0, 1, 2, 3))
    add_body(doc, "Dragon 的方向向量没有再次归一化，因此外侧炸弹的实际合速度高于中间炸弹。垂直俯冲使用 sin(timer × 1.5)，水平摆动频率由移动速度 80 决定。")

    add_heading(doc, "5.2 Monster 攻击细节", 2)
    add_body(doc, "Monster 每轮先计算从 Boss 到玩家的单位方向 (dx, dy)，然后生成三颗炸弹；三颗弹的 X 出生偏移分别为 -40、0、+40，方向 X 分量再分别加 -0.2、0、+0.2。子弹发射后不会持续修正方向，因此属于“发射瞬间瞄准”，不是持续追踪。")
    add_callout(doc, "已知攻击计时问题", "两个 Boss 的 update() 先执行 Enemy::update()，随后又对同一个 bombTimer 增加 dt，并另外维护 Boss 弹幕；父类还可能创建一颗未被 Boss 绘制/碰撞处理的普通炸弹。当前有效攻击节奏可能受到双重计时影响，建议重构后重新测定。", RISK_FILL)

    add_heading(doc, "6. 道具与危险物", 1)
    add_table(doc, ["类型", "外观资源", "效果", "下落速度", "单次生成概率"], [
        ("HealthPickup", "things_silver.png", "玩家生命 +1，最大为 3", "150 像素/秒", "25%"),
        ("ScorePickup", "star_gold.png", "分数 +50", "150 像素/秒", "25%"),
        ("PiercingPickup", "powerupGreen_bolt.png", "5 秒穿透子弹", "150 像素/秒", "25%"),
        ("MeteorHazard", "4 种随机陨石", "玩家生命 -1", "250 像素/秒", "25%"),
    ], [1300, 2200, 2800, 1600, 1460], center_cols=(0, 3, 4))
    add_table(doc, ["生成参数", "当前值", "推导"], [
        ("检查间隔", "每 5 秒", "pickupSpawnTimer > 5.0"),
        ("每次检查生成概率", "40%", "rand()%100 < 40"),
        ("每次检查不生成概率", "60%", "未通过概率判断"),
        ("长期平均生成间隔", "约 12.5 秒/个", "5 秒 ÷ 40%"),
        ("单种物品每次检查概率", "10%", "40% × 25%"),
        ("随机 X 范围", "0~949", "rand() % (1000-50)"),
        ("初始 Y", "-50", "从屏幕上方进入"),
        ("离屏销毁", "y > 1000", "到达底部后销毁"),
        ("MeteorHazard 旋转速度", "-100~+99 度/秒", "rand()%200 - 100"),
    ], [2400, 2400, 4560], center_cols=(0, 1))

    add_heading(doc, "7. 关卡参数", 1)
    add_body(doc, "普通关卡中的每个敌人会独立、等概率地随机成为 Alpha、Beta 或 Gamma，因此每次新游戏的敌人构成和总血量都可能不同。")
    add_table(doc, ["显示等级", "状态", "阵型/敌人", "数量", "位置与范围", "总 HP 范围", "敌人分数范围"], [
        ("1", "Level1", "3×5 网格；A/B/G 随机", "15", "x=140/320/500/680/860；y=50/190/330", "45~105", "150~450"),
        ("2", "Level2", "圆形阵型；A/B/G 随机", "12", "中心 (500,280)，半径 250", "36~84", "120~360"),
        ("3", "Level3", "4×5 网格；A/B/G 随机", "20", "同 L1 的 x；y 再增加 470", "60~140", "200~600"),
        ("4", "BossDragon", "Dragon", "1", "初始 (500,200)", "100", "100"),
        ("5", "BossMonster", "Monster", "1", "初始 (500,200)", "100", "100"),
        ("0", "Victory", "无敌人", "0", "当前立即循环回 Level1", "0", "0"),
    ], [750, 1350, 2150, 700, 2200, 1000, 1210], center_cols=(0, 1, 3, 5, 6))

    add_heading(doc, "7.1 随机敌人构成的期望值", 2)
    add_table(doc, ["关卡", "每种敌人期望数量", "期望总 HP", "期望敌人分数"], [
        ("Level 1（15只）", "各 5 只", "75", "300"),
        ("Level 2（12只）", "各 4 只", "60", "240"),
        ("Level 3（20只）", "各约 6.67 只", "100", "400"),
        ("普通关卡合计", "各约 15.67 只", "235", "940"),
        ("含两个 Boss", "另加 Dragon + Monster", "435", "1140"),
    ], [2300, 2700, 2000, 2360], center_cols=(0, 1, 2, 3))
    add_callout(doc, "理论总分", "只计算击杀敌人时，完整一轮最低 670 分、最高 1610 分、期望约 1140 分；ScorePickup 每次另加 50 分。", GREEN_FILL)

    add_heading(doc, "8. 画面、动画与 HUD", 1)
    add_table(doc, ["项目", "当前参数", "说明"], [
        ("背景缩放", "2.0 × 1.5", "background_1.png"),
        ("爆炸帧数", "4 帧", "横向精灵表"),
        ("爆炸单帧尺寸", "138 × 138", "每帧截取区域"),
        ("爆炸单帧时长", "0.1 秒", "总时长约 0.4 秒"),
        ("爆炸缩放", "0.5", "击杀敌人时生成"),
        ("生命条背景", "200 × 20", "位置 (20,20)"),
        ("生命条比例", "health / 3", "超过 3 HP 时会超过背景宽度"),
        ("生命条颜色", ">50% 绿色，否则红色", "按 health/3 判断"),
        ("HUD 字号", "24", "分数和关卡"),
    ], [2300, 2300, 4760], center_cols=(0, 1))

    add_heading(doc, "9. 计分、排行榜与存档", 1)
    add_table(doc, ["功能", "当前规则", "文件/限制"], [
        ("敌人计分", "A 10 / B 20 / G 30 / Dragon 100 / Monster 100", "敌人死亡时增加"),
        ("道具计分", "ScorePickup +50", "拾取时增加"),
        ("排行榜", "按分数降序，只保留前 10 名", "assets/scores.txt"),
        ("玩家名", "最多 15 个 ASCII 字符", "不支持中文输入；空名字不能提交"),
        ("自动存档", "暂停后返回主菜单时保存", "assets/savegame.bin"),
        ("新游戏", "删除旧 savegame.bin", "然后重置到 Level 1"),
        ("继续游戏", "仅在存档文件存在时显示", "读取失败没有详细用户提示"),
    ], [1900, 4300, 3160], center_cols=(0,))

    add_heading(doc, "9.1 当前存档字段", 2)
    add_table(doc, ["顺序", "字段", "类型/内容"], [
        ("1", "level", "int，当前等级 1~5"),
        ("2", "score", "int，玩家分数"),
        ("3", "health", "int，玩家生命"),
        ("4", "selectedShipType", "int，1~3"),
        ("5", "mouseCtrl", "int，0/1"),
        ("6", "playerX / playerY", "两个 float"),
        ("7", "enemyCount", "int"),
        ("8", "每个敌人", "EnemyType（1 字节 ASCII 代码）+ float x/y + int HP"),
    ], [900, 2800, 5660], center_cols=(0, 1))
    add_callout(doc, "未保存状态", "玩家子弹、敌人炸弹、拾取物、玩家速度、穿透效果剩余时间、攻击计时器、爆炸动画和 Boss 运动计时均不在存档中。", NOTE_FILL)

    add_heading(doc, "10. 当前实现中应优先统一的参数", 1)
    add_table(doc, ["优先级", "问题", "建议"], [
        ("高", "生命上限同时存在 3 和 10", "建立统一 PlayerConfig::maxHealth，并让 HUD、道具、存档共用"),
        ("高", "Boss 和父类重复操作 bombTimer", "把攻击更新改为可覆盖的虚函数或独立 AttackPattern"),
        ("高", "Victory 立即循环到 Level 1", "真正进入 Victory 状态并展示胜利界面"),
        ("中", "Boss 攻击向量未统一归一化", "明确要恒定合速度还是恒定垂直速度"),
        ("中", "随机数使用 std::rand", "改用 <random> 与可注入种子，便于测试和复现"),
    ], [1000, 3700, 4660], center_cols=(0,))

    add_heading(doc, "11. 参数修改索引", 1)
    add_table(doc, ["要修改的内容", "文件"], [
        ("窗口与世界尺寸", "src/config/GameConfig.hpp"),
        ("帧率、爆炸参数", "src/app/Game.hpp、Game.cpp"),
        ("玩家生命、速度、加速度、摩擦、射击冷却", "src/entities/Spaceship.hpp、Spaceship.cpp"),
        ("玩家子弹速度与销毁边界", "src/entities/Bullet.hpp、Bullet.cpp"),
        ("普通敌人 HP、弹速、冷却、移动速度", "src/entities/enemies/InvaderTypes.hpp、InvaderTypes.cpp"),
        ("普通敌人公共摆动与攻击逻辑", "src/entities/enemies/Enemy.hpp、Enemy.cpp"),
        ("Dragon 参数", "src/entities/enemies/Dragon.hpp、Dragon.cpp"),
        ("Monster 参数", "src/entities/enemies/Monster.hpp、Monster.cpp"),
        ("敌人炸弹基础参数", "src/entities/enemies/Bomb.hpp、Bomb.cpp"),
        ("关卡数量、阵型、坐标", "src/gameplay/LevelManager.hpp、LevelManager.cpp"),
        ("道具效果和速度", "src/entities/pickups/*.hpp"),
        ("道具生成周期与概率", "src/app/Game.cpp：spawnPickup/update"),
        ("排行榜数据与持久化", "src/ui/Leaderboard.hpp、Leaderboard.cpp"),
        ("排行榜界面与其他菜单", "src/ui/Menu.hpp、Menu*.cpp"),
        ("存档格式", "src/systems/SaveSystem.cpp；src/app/Game.cpp：对象还原"),
    ], [3600, 5760], center_cols=(0,))

    add_callout(doc, "维护建议", "每次调整游戏平衡后同步更新本手册，并重新构建运行；更理想的做法是把这些数值迁移到集中式配置或 JSON 文件，让程序和文档从同一份数据生成。", GREEN_FILL)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Nova Strike 参数手册"
    doc.core_properties.subject = "当前源码中的玩家、敌人、道具、关卡与系统参数"
    doc.core_properties.keywords = "Nova Strike, C++, SFML, 参数, 游戏平衡"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
